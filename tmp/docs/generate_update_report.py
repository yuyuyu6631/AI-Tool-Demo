from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\codespace\workfile\output\doc\星点评-网站迭代更新报告-2026-04-22.docx")
DESKTOP_SCREENSHOT = Path(r"D:\codespace\workfile\apps\web\src\e2e\visual.spec.ts-snapshots\home-1920-win32.png")
MOBILE_SCREENSHOT = Path(r"D:\codespace\workfile\apps\web\src\e2e\visual.spec.ts-snapshots\home-search-390-win32.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_style(document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(document, text, bold=False, color=None, size=None, align=None, space_after=6):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_screenshot_section(document):
    document.add_heading("当前网站样式概览", level=1)
    add_paragraph(
        document,
        "以下截图基于仓库内最近一次前端视觉回归快照整理，用于展示当前网站首页在桌面端与移动端的整体视觉样式。",
    )

    if DESKTOP_SCREENSHOT.exists():
        add_paragraph(document, "图 1：桌面端首页样式", bold=True, color="1F4E78", space_after=4)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(DESKTOP_SCREENSHOT), width=Inches(6.2))
        add_paragraph(document, "桌面端页面已形成统一首屏结构，品牌区、搜索入口、分类与内容卡片层次比较完整。", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    if MOBILE_SCREENSHOT.exists():
        add_paragraph(document, "图 2：移动端首页样式", bold=True, color="1F4E78", space_after=4)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(MOBILE_SCREENSHOT), width=Inches(3.0))
        add_paragraph(document, "移动端保留了首页核心搜索与内容浏览能力，整体已经具备较好的信息聚合和首屏可读性。", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def add_milestone_table(document):
    document.add_heading("一、阶段里程碑总览", level=1)
    add_paragraph(
        document,
        "本报告基于 2026 年 4 月 22 日已提交代码基线整理，覆盖网站本轮从 MVP 补强到内容运营后台成型的主要迭代成果。当前工作区仍存在少量未提交改动，以下内容不将其计入“已完成”范围。",
    )

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["日期", "阶段", "核心更新", "业务意义"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "DCE6F1")

    rows = [
        [
            "2026-04-01",
            "MVP 第一阶段收口",
            "完成排行榜、场景导航、评分展示、价格筛选，并将目录筛选/排序/分页下推到 SQLAlchemy 查询层。",
            "解决演示期核心体验缺口，显著改善工具目录页的性能和可用性。",
        ],
        [
            "2026-04-08",
            "搜索底座升级",
            "完成 Phase 1 embedding recall，引入 tool_embeddings 持久层、回填脚本与回退机制，并补齐验收文档。",
            "搜索结果从纯词法匹配升级为“词法 + 语义召回”，为后续 AI 搜索和推荐打底。",
        ],
        [
            "2026-04-14",
            "AI 搜索与对比体验接入",
            "新增 /api/ai-search、RAG chat、工具对比页、命令面板与首页统一搜索入口，补齐聊天和对比测试。",
            "网站从静态目录升级为可理解用户意图、支持对比决策的 AI 工具发现平台。",
        ],
        [
            "2026-04-22",
            "后台评测工作流与目录体验刷新",
            "新增后台总览、工具管理、榜单管理、评测审核接口与页面，补齐评论管理、首页聚合数据、文档同步基线。",
            "平台开始具备“前台发现 + 后台运营 + 评测治理”的闭环能力。",
        ],
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text


def add_completed_updates(document):
    document.add_heading("二、已完成更新内容说明", level=1)

    sections = [
        (
            "1. 产品与信息架构",
            [
                "前台已形成统一首页入口，并稳定承载直接搜索、AI 帮找、榜单、场景、工具详情与认证/登录入口。",
                "工具详情页已补齐评分、价格、标签、相关推荐和用户评测读取能力，用户决策链路更完整。",
                "目录页与首页聚合能力增强后，网站当前更接近“AI 工具版大众点评/豆瓣”的产品形态，而不再只是静态收录站。",
            ],
        ),
        (
            "2. AI 能力与搜索底座",
            [
                "搜索接口已升级为词法检索与 embedding 召回混合方案，在 embedding 不可用时可无感回退到原有词法搜索。",
                "新增 AI Search 与 RAG Chat 后，用户可以通过自然语言描述需求、场景或目标，由系统给出工具建议与解释。",
                "工具解析服务和推荐链路已经接入，为后续自动补录、专题编排和更强推荐提供了服务层基础。",
            ],
        ),
        (
            "3. 后台运营与内容治理",
            [
                "后台已形成 overview、tools、rankings、reviews 四类管理入口，具备工具新增/编辑、榜单维护和评测审核能力。",
                "评论相关接口已支持读取、提交、查询我的评论与后台删除，基础审核闭环已经具备。",
                "数据库模型已扩展到用户、会话、评测、更新提案、榜单和 embedding 等关键实体，支撑持续运营。",
            ],
        ),
        (
            "4. 工程质量与交付可维护性",
            [
                "当前前端共 15 个页面路由，后端共 34 个 API 端点，单仓架构已覆盖展示、搜索、评测、后台和基础认证。",
                "仓库内已有 17 个 API pytest 文件、19 个前端单元/集成测试文件、4 个 E2E 脚本，关键主链路具备回归保护。",
                "文档基线已接入自动同步与 pre-commit 刷新流程，降低“代码变了但说明没跟上”的协作成本。",
            ],
        ),
    ]

    for title, bullets in sections:
        document.add_heading(title, level=2)
        add_bullets(document, bullets)


def add_value_and_expectation(document):
    document.add_heading("三、本轮迭代带来的直接价值", level=1)
    add_bullets(
        document,
        [
            "对用户侧：从“找工具”升级为“理解需求、给出建议、辅助比较与决策”。",
            "对运营侧：从手工散点维护升级为可管理工具、榜单和评测内容的后台工作台。",
            "对项目侧：产品路线已从 MVP 演示站进入可持续扩展的业务基线，具备继续放量建设的条件。",
        ],
    )

    document.add_heading("四、下一步计划", level=1)
    next_steps = [
        ("优先级 P1：评测与数据治理闭环", "补齐抓取结果审核、工具更新提案处理、内容发布状态流转，以及评测质量标准与违规规则。"),
        ("优先级 P1：AI 搜索可解释性与推荐质量", "增加召回来源说明、排序依据、推荐理由结构化输出，并持续优化语义召回命中率。"),
        ("优先级 P2：增长与内容能力", "建设 SEO 专题页、主题榜单模板、场景页内容增强和用户留资/订阅触点。"),
        ("优先级 P2：运营指标与监控", "补齐搜索点击、转化、评测提交、榜单点击等埋点，并建立错误告警与接口稳定性监控。"),
        ("优先级 P3：权限与协作能力", "细化管理员、编辑、审核者等角色权限，支撑多人协作运营。"),
    ]
    for title, desc in next_steps:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(f"{title}：")
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
        run2 = p.add_run(desc)
        run2.font.name = "Microsoft YaHei"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run2.font.size = Pt(10.5)

    document.add_heading("五、下一阶段预期", level=1)
    add_bullets(
        document,
        [
            "如果按上述优先级继续推进，下一阶段可以把平台从“可用”推进到“可运营、可放量、可衡量”。",
            "预计将显著提升搜索命中感知、工具页转化效率、榜单与专题内容复用率，以及后台运营效率。",
            "一旦内容治理和埋点体系稳定，项目就可以开始进入真实用户验证和增长实验阶段。",
        ],
    )


def add_risk_table(document):
    document.add_heading("六、当前主要风险与建议应对", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["风险项", "表现", "影响", "建议应对"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "FCE4D6")

    risks = [
        ["数据质量风险", "工具资料、评分、标签与嵌入内容可能出现脏数据或更新滞后。", "影响 AI 推荐可信度与搜索准确率。", "建立审核流、回填任务和异常数据巡检脚本。"],
        ["AI 结果可解释性不足", "用户暂时难以理解为什么推荐这些工具，召回来源不透明。", "影响用户信任和复用率。", "在搜索结果和聊天回答中补充推荐理由、来源与排序说明。"],
        ["后台权限仍偏粗", "当前后台能力已具备，但角色边界还不够细。", "多人协作时容易出现误操作或责任不清。", "尽快细化角色模型，并对高风险操作增加审计日志。"],
        ["缺少运营指标闭环", "已有测试和文档基线，但增长与运营指标还未系统化采集。", "难以判断哪些页面和功能真正产生价值。", "优先补齐埋点、仪表盘和周报机制。"],
        ["工作区仍有未提交改动", "当前仓库存在部分服务层与测试文件改动。", "若直接对外汇报为“已完成”，容易与真实已交付边界混淆。", "对外口径统一采用已提交基线，未提交内容单列为“进行中”。"],
    ]

    for row in risks:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text


def add_footer_summary(document):
    document.add_heading("七、汇总结论", level=1)
    add_paragraph(
        document,
        "截至 2026 年 4 月 22 日，星点评网站已经完成从 MVP 展示站到“AI 工具发现 + 评测 + 后台运营”平台基线的关键跨越。当前最重要的不是继续堆新功能，而是把内容治理、AI 推荐解释、运营指标和权限体系补齐，让现有能力真正稳定地产生业务价值。",
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_paragraph(document, "附录：报告依据", bold=True, size=12, color="1F4E78")
    add_bullets(
        document,
        [
            "代码基线日期：2026-04-22",
            "关键提交：efb46ae、a35d343、de56f17、f8645f6",
            "参考材料：README.md、doc/phase1-search-update-summary.md、docs/current-implementation-baseline.md、git log / git show 结果",
        ],
    )


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_page_style(document)

    add_paragraph(document, "星点评网站迭代更新报告", bold=True, size=20, color="1F4E78", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(document, "报告日期：2026 年 4 月 22 日", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_paragraph(document, "适用对象：项目负责人、产品、研发、运营", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_paragraph(
        document,
        "本报告用于总结星点评网站本轮迭代成果，说明已经完成的关键更新、下一步推进计划、阶段性预期，以及当前需要重点关注的风险项。报告内容以仓库已提交代码基线为准。",
        space_after=10,
    )

    add_screenshot_section(document)
    add_milestone_table(document)
    add_completed_updates(document)
    add_value_and_expectation(document)
    add_risk_table(document)
    add_footer_summary(document)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
